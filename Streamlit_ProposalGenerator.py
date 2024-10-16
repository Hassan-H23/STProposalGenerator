from datetime import datetime
from io import BytesIO
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt
import docxedit
import pandas as pd
import streamlit as st
from streamlit import divider, title
import Service
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Inches


st.set_page_config(layout="wide",page_icon="icons/Maverick Single - White.png",page_title="Maverick Proposals")


def initialize_session_state():
    if 'services_list' not in st.session_state:
        st.session_state.services_list = []
def add_bullet(paragraph):

    p = paragraph._element
    pPr = p.get_or_add_pPr()

    # Create bullet properties
    numPr = OxmlElement('w:numPr')  # Create numbering element
    ilvl = OxmlElement('w:ilvl')  # Indentation level
    ilvl.set(qn('w:val'), '0')  # Set indentation level for the bullet
    numId = OxmlElement('w:numId')  # Create list item
    numId.set(qn('w:val'), '1')  # Set numbering ID to correspond to bullets

    # Define bullet character (this can be adjusted to different bullet styles)
    bulletChar = OxmlElement('w:bullet')
    bulletChar.set(qn('w:val'), '•')  # Use the bullet character

    # Append bullet properties to the numPr

    numPr.append(ilvl)  # Add indentation to the numbering properties
    numPr.append(numId)  # Add numbering ID to the properties
    numPr.append(bulletChar)  # Add bullet character

    # Append numPr to paragraph properties
    pPr.append(numPr)  # Append the numbering properties to the paragraph properties


def addScope(document, content):
    header_paragraph = document.add_heading(level=0)
    run = header_paragraph.add_run()
    run.add_picture("icons/Maverick Logo.png", width=Inches(6))
    document.add_paragraph()

    # Iterate through the content to process lines
    for line in content:
        if '#' in line:
            line = line.replace('#', '')
            heading = document.add_heading(level=1)
            run = heading.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
            heading.paragraph_format.line_spacing = 0.85
        elif '*' in line:
            line = line.replace('*', '')
            para = document.add_paragraph(line)
            para.paragraph_format.line_spacing = 0.85
        else:
            para = document.add_paragraph(line)
            para.paragraph_format.line_spacing = 0.85
            add_bullet(para)


    document.add_paragraph()
    scope_paragraph = document.add_paragraph("***Scope of work subject to be adjusted per client's request***")
    scope_paragraph.paragraph_format.alignment = 1  # Center alignment
    scope_paragraph.paragraph_format.line_spacing = 1.0

    document.add_page_break()



def set_cell_margins(cell, top, right, bottom, left):
    cell._element.get_or_add_tcPr().append(
        parse_xml(
            r'<w:tcMar {}>'
            r'<w:top w:w="{top}" w:type="dxa"/>'
            r'<w:right w:w="{right}" w:type="dxa"/>'
            r'<w:bottom w:w="{bottom}" w:type="dxa"/>'
            r'<w:left w:w="{left}" w:type="dxa"/>'
            r'</w:tcMar>'.format(nsdecls('w'), top=top, right=right, bottom=bottom, left=left)
        )
    )

def get_ordinal_suffix(day):
    if 10 <= day % 100 <= 20:
        return "th"
    else:
        return {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")


def read_file(file_name):
    lines = []
    try:
        with open(file_name, 'r') as file:
            for line in file:
                stripped_line = line.strip()
                if stripped_line:
                    lines.append(stripped_line)
    except FileNotFoundError:
        print(f"Error: The file '{file_name}' was not found.")
    except Exception as e:
        print(f"An error occurred: {e}")

    return lines


def populate_set(Set):
    for service in st.session_state.services_list:
        Set.add(service.serviceName)

    # Convert the set to a list to maintain the order
    service_list = list(Set)
    buildings = []

    # Check if there are more than one service to concatenate "and"
    if len(service_list) > 1:
        # Modify the last element to include "and"
        service_list[-1] = f"and {service_list[-1]}"

    # Join the elements with commas except the last one which has "and"
    return ", ".join(service_list)


# Call the function to initialize the session state
initialize_session_state()

if 'Service_Names_Set' not in st.session_state:
    st.session_state.Service_Names_Set = set()

st.image("icons/Maverick_Proposal.png", width=450)
# File uploader for the proposal template
st.header('Service Table', divider="gray")

# Sidebar for company details and service input
with st.sidebar:
    st.logo("icons/Maverick Single - White.png",size="large")
    companyName = st.text_input("Building Name")
    ceoName = st.text_input("Attention To")
    companyAddress = st.text_input("Building Address")

    with st.form(key="Services"):
        serviceChoice = st.selectbox("Service", ["Concierge", "Security", "Janitorial", "Cleaning", "Porter", "Valet"])
        weeklyHours = st.number_input("Weekly Hours", min_value=0, max_value=1000, step=1, value=168)
        billRate = st.number_input("Bill Rate", min_value=0.0, max_value=1000.0, step=1.0, value=27.0)
        yearlyHolidayHours = st.number_input("Yearly Holiday Hours", min_value=0, max_value=1000, step=1)
        inflationRate = st.number_input("Inflation Rate", min_value=0.0, max_value=100.0, step=0.1, value=3.0)

        submitted = st.form_submit_button("Add Service")
        finalSubmit = st.form_submit_button("Generate Proposal")

    if submitted:
        new_service = Service.Service(serviceChoice, weeklyHours, billRate, yearlyHolidayHours, inflationRate,0.0,0.0,0.0,0.0)
        st.session_state.services_list.append(new_service)
        st.success(f"{serviceChoice} service added successfully!")
        st.rerun()

    if finalSubmit:
        # Validation

        if not companyName or not ceoName or not companyAddress:
            st.warning("Please fill in all company details: Company Name, CEO Name, and Company Address.")
        elif not st.session_state.services_list:
            st.warning("Please add at least one service before generating the proposal.")
        else:

            document = Document("ScopesOfWork/Maverick Proposal Template .docx")
            ScopeOfWork = populate_set(st.session_state.Service_Names_Set)
            current_date = datetime.now()
            formatted_date = current_date.strftime(
                "%B") + f" {current_date.day}{get_ordinal_suffix(current_date.day)}, {current_date.year}"

            docxedit.replace_string(document, 'CCNN', companyName)
            docxedit.replace_string(document, 'CCAA', str(companyAddress))
            docxedit.replace_string(document, 'NNNN', ceoName)
            docxedit.replace_string(document, 'DDDD', formatted_date)
            docxedit.replace_string(document, 'SSSS', ScopeOfWork)

            # Scope of Work (Jan and Conc Only)
            for service in st.session_state.services_list:
                if service.serviceName == "Janitorial":
                    content = read_file("ScopesOfWork/JANITORIAL.txt")
                    addScope(document, content)
                if service.serviceName == "Concierge":
                    content = read_file("ScopesOfWork/CONCIERGE.txt")
                    addScope(document, content)

            # Cost Proposal Table


            document.add_paragraph()
            heading = document.add_heading('COST PROPOSAL', level=1)
            run = heading.runs[0]
            run.bold = True
            run.font.size = Pt(22)

            table = document.add_table(rows=1, cols=9)
            table.style = 'Table Grid'

            # Table Headers
            hdr_cells = table.rows[0].cells
            table_headers = [
                "Service", "Weekly Hours", "Annual Holiday Hours", "Bill Rate", "Inflation Rate", "Monthly Amount",
                "Annual Amount (Year 1)", "Annual Amount (Year 2)", "Annual Amount (Year 3)"
            ]

            # Set headers
            for i, header in enumerate(table_headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                hdr_cells[i].paragraphs[0].runs[0].font.size = Inches(0.15)  # Adjust as necessary
                hdr_cells[i].paragraphs[0].paragraph_format.alignment = 1  # Center alignment
                set_cell_margins(hdr_cells[i], 55, 50, 45, 50)  # Adjust margins

                # Set background color
                if i != 0:
                    hdr_cells[i]._element.get_or_add_tcPr().append(
                        parse_xml(r'<w:shd {} w:fill="ADD8E6"/>'.format(nsdecls('w'))))
                else:
                    hdr_cells[i]._element.get_or_add_tcPr().append(
                        parse_xml(r'<w:shd {} w:fill="000080"/>'.format(nsdecls('w'))))

            # Populate the table with service data
            for service in st.session_state.services_list:
                row_cells = table.add_row().cells
                row_cells[0].text = service.serviceName
                row_cells[1].text = str(service.weeklyHours)
                row_cells[2].text = f"{service.yearlyHolidayHours}"
                row_cells[3].text = f"${service.billRate:}"
                row_cells[4].text = f"{service.inflationRate * 100}%"
                row_cells[5].text = f"${service.monthlyAmount}"
                row_cells[6].text = f"${service.annualAmountYear1}"
                row_cells[7].text = f"${service.annualAmountYear2}"
                row_cells[8].text = f"${service.annualAmountYear3}"

                # Set formatting for each cell in the row
                for cell in row_cells:
                    cell.paragraphs[0].paragraph_format.alignment = 1  # Center alignment
                    set_cell_margins(cell, 100, 100, 100, 100)  # Adjust margins

            # Footer

            document.add_paragraph()
            document.add_paragraph("***Applicable NJ Sales Tax Included***").paragraph_format.alignment = 1
            document.add_paragraph(
                "***New Year’s Day, Presidents Day, Memorial Day, Independence Day, Labor Day, Thanksgiving Day, Christmas Day Is Included In the Above Pricing***")
            # Final Page
            document.add_page_break()
            header_paragraph = document.add_heading(level=0)
            run = header_paragraph.add_run()
            run.add_picture("icons/Maverick Logo.png", width=Inches(6))
            document.add_paragraph()
            document.add_paragraph(formatted_date)
            document.add_paragraph("To whom it may concern,")
            document.add_paragraph("This Agreement will be for a one (1) year period. Any termination by either party or other changes are subject to written notice by either party not less than ninety (90) days prior to the end of the Contract Term.")
            document.add_paragraph("Annual rates are subject to increase after each year, refer to the Cost Proposal table for more details.")
            document.add_paragraph("Holiday differential is not included.")
            document.add_paragraph("On behalf of our entire company staff, I would like to express our gratitude for your consideration in this RFP. We are committed to providing effective cooperation and assurance that we will work diligently towards a successful outcome. Thank you for the opportunity.")
            document.add_paragraph()
            document.add_paragraph("Sincerely,")

            document.add_picture("icons/sign.png", width=Inches(1.5))
            document.add_paragraph("Mark Morcos")
            document.add_paragraph("Founder/President")
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            heading = document.add_heading('COMMITMENT TO PROJECT', level=1)
            run = heading.runs[0]
            run.bold = True
            run.font.size = Pt(22)
            document.add_paragraph("We guarantee the association that our team will remain with the project through the entire duration of the contract. Our team is committed to providing reliable, high-quality services, and we understand the importance of consistency and continuity for our clients. We will ensure that our team is fully staffed and that we have the necessary resources to deliver exceptional service throughout the entire contract period. If for any reason there are changes to the team, we will notify the association in advance and take steps to ensure a smooth transition of service")

            doc_io = BytesIO()
            document.save(doc_io)
            doc_io.seek(0)

            # Download button for the proposal
            if st.session_state.services_list:
                st.download_button(
                    label="Download Proposal",
                    data=doc_io,
                    file_name=f"{companyName} Proposal.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# Convert the list of services to a DataFrame

if st.session_state.services_list:
    services_df = pd.DataFrame([vars(service) for service in st.session_state.services_list])

    # Create a mapping of old column names to new header names
    header_mapping = {
        'serviceName': 'Service Name',
        'weeklyHours': 'Weekly Hours',
        'billRate': 'Bill Rate',
        'yearlyHolidayHours': 'Yearly Holiday Hours',
        'inflationRate': 'Inflation Rate (%)',
        'monthlyAmount': 'Monthly Amount',
        'annualAmountYear1': 'Annual Amount Year 1',
        'annualAmountYear2': 'Annual Amount Year 2',
        'annualAmountYear3': 'Annual Amount Year 3'
    }

    # Rename the DataFrame columns
    services_df.rename(columns=header_mapping, inplace=True)

    # Format the values
    services_df['Bill Rate'] = services_df['Bill Rate'].apply(lambda x: f"${x:.2f}")
    services_df['Yearly Holiday Hours'] = services_df['Yearly Holiday Hours'].apply(lambda x: f"{x:.2f}")
    services_df['Inflation Rate (%)'] = services_df['Inflation Rate (%)'].apply(lambda x: f"{int(x * 100)}%")  # Convert to whole number and add '%' sign
    services_df['Monthly Amount'] = services_df['Monthly Amount'].apply(lambda x: f"${x:.2f}")
    services_df['Annual Amount Year 1'] = services_df['Annual Amount Year 1'].apply(lambda x: f"${x:.2f}")
    services_df['Annual Amount Year 2'] = services_df['Annual Amount Year 2'].apply(lambda x: f"${x:.2f}")
    services_df['Annual Amount Year 3'] = services_df['Annual Amount Year 3'].apply(lambda x: f"${x:.2f}")

    # Add ID column based on index
    services_df.insert(0, 'ID', range(1, len(services_df) + 1))

    # Convert the DataFrame to an HTML table with specific widths
    html_table = services_df.to_html(classes='styled-table', index=False, escape=False)

    # Center the table and style it
    st.markdown(
        """
        <style>
        .styled-table {
            width: 100%;
            max-width: 100%; /* Maximum width to fit any screen */
            border-collapse: collapse;
            table-layout: auto; /* Allow automatic column width */
            margin: auto; /* Center the table */
        }
        .styled-table th, .styled-table td {
            text-align: center; /* Center content in cells */
            padding: 12px; /* Add padding for better spacing */
            border: 0.5px solid #ddd; /* Add borders to cells */
            overflow: hidden; /* Prevent overflow */
            text-overflow: ellipsis; /* Ellipsis for overflowed text */
            word-wrap: break-word; /* Break long words to fit */
        }
        .styled-table th {
            background-color: #262730; /* Set header background color */
            color: white; /* Set header text color to white */
            font-weight: bold; /* Make header text bold */
            font-size: 0.95em; /* Decrease font size for headers */
        }
        .styled-table td {
            font-size: 0.9em; /* Decrease font size for data cells */
        }
        .styled-table td:first-child { /* Apply styles specifically for the ID column */
            width: 30px; /* Set a fixed width for the ID column */
        }
        /* Responsive styling */
        @media (max-width: 600px) {
            .styled-table th, .styled-table td {
                padding: 8px; /* Reduce padding on smaller screens */
                font-size: 0.7em; /* Further adjust font size for smaller screens */
            }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Insert an additional <div> to wrap the HTML table for easier control over display
    st.markdown(f"<div style='overflow-x:auto;'>{html_table}</div>", unsafe_allow_html=True)

    # Add custom CSS for the selectbox
    st.markdown(
        """
        <style>
        .custom-select {
            width: 150px; /* Set the desired width */
            font-size: 0.8em; /* Decrease the font size */
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    # Allow user to select a service
    service_to_edit = st.selectbox(
        "Select Service to Remove",
        options=[f"{index + 1} - {service.serviceName}" for index, service in
                 enumerate(st.session_state.services_list)] + ["None"],
        key="service_select",
        index=0,
    )
    if service_to_edit != "None":
        selected_service_index = int(service_to_edit.split(" - ")[0]) - 1  # Adjust index back to 0-based for selection
        selected_service = st.session_state.services_list[selected_service_index]

        if st.button("Remove Selected Service"):
             # Remove the service from the list
            st.session_state.services_list.pop(selected_service_index)
            st.success("Service removed successfully!")
            st.rerun()
else:
    st.write("No services added")
